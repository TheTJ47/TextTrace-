# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import requests
# import io
# import json
# from fpdf import FPDF
# import tempfile

# app = Flask(__name__)

# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR on the image
#     url_api = "https://api.ocr.space/parse/image"
#     _, compressedimage = cv2.imencode(".jpg", img, [1, 90])
#     file_bytes = io.BytesIO(compressedimage)

#     result = requests.post(url_api,
#                            files={"screenshot.jpg": file_bytes},
#                            data={"apikey": "helloworld",
#                                  "language": "eng"})

#     if result.status_code == 200:
#         result = result.content.decode()
#         result = json.loads(result)
#         parsed_results = result.get("ParsedResults", [])  # Get ParsedResults or an empty list if it doesn't exist
#         if parsed_results:
#             text_detected = parsed_results[0].get("ParsedText")

#             # Remove non-ASCII characters from the text
#             text_detected_cleaned = ''.join(char for char in text_detected if ord(char) < 128)

#             # Convert detected text to PDF
#             pdf = FPDF()
#             pdf.add_page()
#             pdf.set_font("Arial", size=12)
#             pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#             pdf.cell(200, 10, txt="", ln=True)
#             pdf.multi_cell(0, 10, txt=text_detected_cleaned)

#             # Create a temporary file to save the PDF
#             with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#                 pdf.output(tmp_file.name)
#                 tmp_file.seek(0)

#             # Send the temporary file as an attachment
#             response = make_response(send_file(tmp_file.name))
#             response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#             return response
#         else:
#             return jsonify({'error': 'No text detected in the image'}), 400
#     else:
#         return jsonify({'error': f'OCR request failed with status code: {result.status_code}'}), result.status_code

# if __name__ == '__main__':
#     app.run(debug=True)

# # from flask import Flask, render_template, request, jsonify, send_file, make_response
# # import cv2
# # import numpy as np
# # from google.cloud import vision  # Import Google Cloud Vision library

# # app = Flask(__name__)

# # # Replace with your Google Cloud project credentials (refer to setup instructions)
# # project_id = "your-project-id"
# # client = vision.ImageAnnotatorClient(client_options={"api_endpoint": "vision.googleapis.com"})

# # @app.route('/')
# # def index():
# #     return render_template('index.html')

# # @app.route('/recognize_handwriting', methods=['POST'])
# # def recognize_handwriting():
# #     # Get the image file from the request
# #     image_file = request.files['image']

# #     # Convert image data to a byte array
# #     image_data = np.frombuffer(image_file.read(), dtype=np.uint8)

# #     # Prepare image as a content type
# #     content = vision.types.Image(content=image_data.tobytes())

# #     # Request text detection features
# #     features = [vision.types.Feature(
# #         type_=vision.types.Feature.Type.DOCUMENT_TEXT_DETECTION)]
    
# #     # Send request to Google Cloud Vision API
# #     response = client.annotate_image(image=content, features=features)

# #     # Extract recognized text
# #     text = response.full_text_annotation.text
    
# #     if text:
# #         # Convert detected text to PDF
# #         pdf = FPDF()
# #         pdf.add_page()
# #         pdf.set_font("Arial", size=12)
# #         pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
# #         pdf.cell(200, 10, txt="", ln=True)
# #         pdf.multi_cell(0, 10, txt=text)

# #         # Create a temporary file to save the PDF
# #         with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
# #             pdf.output(tmp_file.name)
# #             tmp_file.seek(0)

# #         # Send the temporary file as an attachment
# #         response = make_response(send_file(tmp_file.name))
# #         response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
# #         return response
# #     else:
# #         return jsonify({'error': 'No text detected in the image'}), 400

# # if __name__ == '__main__':
# #     app.run(debug=True)

                #TESSERACT OCR CODE

# app.py
# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# import pytesseract

# app = Flask(__name__)

# # Configure pytesseract path to Tesseract executable
# # Update the path below to where Tesseract is installed on your system
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# # Define paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using pytesseract
#     text_detected = pytesseract.image_to_string(img)

#     # Check if text was detected
#     if not text_detected.strip():
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Remove non-ASCII characters from the text
#     text_detected_cleaned = ''.join(char for char in text_detected if ord(char) < 128)

#     # Convert detected text to PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#     pdf.cell(200, 10, txt="", ln=True)
#     pdf.multi_cell(0, 10, txt=text_detected_cleaned)

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         pdf.output(tmp_file.name)
#         tmp_file.seek(0)

#     # Send the temporary file as an attachment
#     response = make_response(send_file(tmp_file.name, as_attachment=True))
#     response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#     return response

# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)

#  #Current Running

# from easyocr import Reader
# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas

# app = Flask(__name__)

# # Configure paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# # Initialize EasyOCR reader (replace 'en' with your required language(s))
# reader = Reader(['en'])

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']
#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using EasyOCR
#     text_detected = reader.readtext(img)

#     # Check if text was detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Process detected text (e.g., join lines, remove non-ASCII characters if not needed)
#     text_detected_cleaned = '\n'.join([line[1] for line in text_detected])

#     # Create PDF using reportlab
#     buffer = io.BytesIO()
#     c = canvas.Canvas(buffer, pagesize=letter)
#     c.drawString(100, 750, "Handwriting Recognition Result")
#     c.drawString(100, 730, text_detected_cleaned)
#     c.save()

#     buffer.seek(0)

#     # Send the PDF as an attachment
#     response = make_response(send_file(buffer, as_attachment=True, attachment_filename="recognized_text.pdf"))

#     return response

# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)


# from easyocr import Reader
# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# import pdfkit

# app = Flask(__name__)

# # Set path to wkhtmltopdf executable
# app.config['WKHTMLTOPDF_PATH'] = 'C:\Program Files\wkhtmltopdf'  # Replace '/path/to/wkhtmltopdf' with the actual path

# # Initialize EasyOCR reader (replace 'en' with your required language(s))
# reader = Reader(['en'])


# # Configure paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# # Initialize EasyOCR reader (replace 'en' with your required language(s))
# reader = Reader(['en'])

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']
#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using EasyOCR
#     text_detected = reader.readtext(img)

#     # Check if text was detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Process detected text (e.g., join lines, remove non-ASCII characters if not needed)
#     text_detected_cleaned = '\n'.join([line[1] for line in text_detected])

#     # HTML template for PDF content
#     html_content = f"""
#     <!DOCTYPE html>
#     <html>
#     <head>
#         <title>Handwriting Recognition Result</title>
#         <style>
#             /* Your CSS styles can be placed here */
#         </style>
#     </head>
#     <body>
#         <h1>Handwriting Recognition Result</h1>
#         <p>{text_detected_cleaned}</p>
#     </body>
#     </html>
#     """

#     # Generate PDF from HTML content
#     pdf = pdfkit.from_string(html_content, False)

#     # Send the PDF as an attachment
#     response = make_response(pdf)
#     response.headers['Content-Type'] = 'application/pdf'
#     response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'

#     return response

# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)

#                         
# import easyocr
# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# from fpdf import FPDF
# import tempfile

# app = Flask(__name__)

# # Initialize EasyOCR reader with desired languages
# reader = easyocr.Reader(['ch_sim', 'en'])

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Check if a file was uploaded
#     if 'image' not in request.files:
#         return jsonify({'error': 'No file part in the request'}), 400

#     # Get the image file from the request
#     image_file = request.files['image']

#     # Decode image data using OpenCV with format detection
#     image_data = image_file.read()
#     img = cv2.imdecode(np.frombuffer(image_data, dtype=np.uint8), cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using EasyOCR
#     text_detected = reader.readtext(img)

#     # Check if text was detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Process detected text (e.g., join lines, remove non-ASCII characters if needed)
#     text_detected_cleaned = ' '.join([line[1] for line in text_detected])  # Join words with spaces

#     # Convert detected text to PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#     pdf.cell(200, 10, txt="", ln=True)

#     # Use utf-8 encoding when writing text to the PDF
#     pdf.multi_cell(0, 10, txt=text_detected_cleaned.encode('latin-1', 'replace').decode('latin-1'))

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         name = "output_file.pdf"
#         pdf.output(name)

#         tmp_file.seek(0)

#         # Send the temporary file as an attachment
#         response = make_response(send_file(tmp_file.name, as_attachment=True))
#         response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#         return response

# if __name__ == '__main__':
#     # Run Flask app

#     app.run(host='0.0.0.0', port=5000, debug=True)








#         # PADDLEOCR
# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# from paddleocr import PaddleOCR

# # Initialize PaddleOCR reader (replace 'en' with your required language(s))
# ocr = PaddleOCR(lang='en')  # Replace 'en' with your desired language code(s)

# # Define paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# app = Flask(__name__)

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using PaddleOCR
#     result = ocr.ocr(img)

#     # Print OCR results for debugging
#     print("OCR Results:")
#     print(result)

#     # Check if text was detected
#     if not result:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     text_detected = []
#     for line in result:
#         for item in line:
#             if isinstance(item, list): # Check if the item is a list (bounding box)
#                 continue
#             text_detected.append(item[0]) # Extract the text part

# # Print detected text for debugging
#     print("Detected Text:")
#     print(text_detected)

# # Check if text was successfully detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

# # Now text_detected is a list of strings, which can be joined
#     text_detected_cleaned = ' '.join(text_detected) # Join all lines with spaces

# # Print cleaned text for debugging
#     print("Text Detected Cleaned:")
#     print(text_detected_cleaned)


#     # Convert detected text to PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#     pdf.cell(200, 10, txt="", ln=True)
#     pdf.multi_cell(0, 10, txt=text_detected_cleaned)

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         pdf.output(tmp_file.name)
#         tmp_file.seek(0)

#     # Print PDF file size for debugging
#     import os
#     print("PDF File Size:", os.path.getsize(tmp_file.name))

#     # Send the temporary file as an attachment
#     response = make_response(send_file(tmp_file.name, as_attachment=True))
#     response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#     return response

# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=500)





#                         # HANDPRINT
# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# import handprint




# app = Flask(__name__)

# # Configure paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500
#     # Perform handwriting recognition using Handprint
#     text_detected = handprint.recognize_image(img)


#     # Check if text was detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Process detected text (e.g., join lines, remove non-ASCII characters if not needed)
#     text_detected_cleaned = '\n'.join(text_detected)

#     # Create PDF using reportlab
#     buffer = io.BytesIO()
#     c = canvas.Canvas(buffer, pagesize=letter)
#     c.drawString(100, 750, "Handwriting Recognition Result")
#     c.drawString(100, 730, text_detected_cleaned)
#     c.save()

#     buffer.seek(0)

#     # Send the PDF as an attachment
#     response = make_response(send_file(buffer, as_attachment=True, attachment_filename="recognized_text.pdf"))

#     return response

# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)


#                                   Tenserflow

# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# import tensorflow as tf

# # Function to preprocess image for TensorFlow model
# def preprocess_image(img, target_size=(224, 224)):  # Adjust based on model requirements
#     """
#     Preprocesses an image for TensorFlow model input.

#     Args:
#         img (np.ndarray): The image to preprocess.
#         target_size (tuple, optional): The target size for resizing. Defaults to (224, 224).

#     Returns:
#         np.ndarray: The preprocessed image.
#     """

#     # Convert to grayscale if needed (adjust based on model)
#     if len(img.shape) == 3 and img.shape[2] == 3:
#         img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

#     # Resize
#     img = cv2.resize(img, target_size)

#     # Normalize (adjust based on model's preprocessing steps)
#     img = img.astype('float32') / 255.0

#     # Expand dimensions if model expects 4D input
#     if len(img.shape) == 2:
#         img = np.expand_dims(img, axis=-1)

#     return img

# app = Flask(__name__)

# # Configure paths to CSS and JavaScript files (unchanged)
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# # Load a pre-trained TensorFlow model for text recognition
# model = tf.keras.models.load_model('path/to/your/model.h5')  # Replace with your model path

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request (unchanged)
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format (unchanged)
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully (unchanged)
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Preprocess the image for TensorFlow model
#     img_preprocessed = preprocess_image(img)

#     # Perform text recognition using TensorFlow
#     predictions = model.predict(np.expand_dims(img_preprocessed, axis=0))  # Add batch dimension

#     # Process detected text based on model output format
#     if isinstance(predictions, np.ndarray):  # Assuming model outputs probabilities
#         # Use a threshold (e.g., 0.5) to identify confident predictions
#         text_detected = [chr(np.argmax(p)) for p in predictions[0] if np.max(p) > 0.5]
#         text_detected_cleaned = ''.join(text_detected)
#     else:
#         # Handle other output formats (e.g., direct text or bounding boxes)
#         # ... (Implement processing based on model output)
#         raise NotImplementedError("Unsupported model output format")

#     # Rest of the code for PDF generation and response remains unchanged
#     pdf = FPDF()
#     # ... (PDF generation logic)

#     # Create a temporary file to save the PDF (unchanged)
#     # ... (Temporary file handling)

#     # Send the temporary file as an attachment (unchanged)
#     response = make_response(send_file(tmp_file.name, as_attachment=True))
#     response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#     return response

# if __name__ == '__main__':
#     # Run Flask app on all network interfaces (unchanged)
#     app.run(host='0.0.0.0', port=5000, debug=True)




                                                #Easy Ocr


# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# from easyocr import Reader
# import os
# os.environ['CUDA_VISIBLE_DEVICES'] = '-1'


# app = Flask(__name__)

# # Configure paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# # Initialize EasyOCR reader (replace 'en' with your required language(s))
# reader = Reader(['en'])

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using EasyOCR
#     text_detected = reader.readtext(img)

#     # Check if text was detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Process detected text (e.g., join lines, remove non-ASCII characters if needed)
#     text_detected_cleaned = ' '.join([line[1] for line in text_detected])  # Join words with spaces

#     # Convert detected text to PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#     pdf.cell(200, 10, txt="", ln=True)
#     pdf.multi_cell(0, 10, txt=text_detected_cleaned)

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         pdf.output(tmp_file.name)
#         tmp_file.seek(0)

#     # Send the temporary file as an attachment
#     response = make_response(send_file(tmp_file.name, as_attachment=True))
#     response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#     return response


# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)





# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# from easyocr import Reader
# import os
# import docx  # Import the python-docx library

# os.environ['CUDA_VISIBLE_DEVICES'] = '-1'
# app = Flask(__name__)

# # Configure paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# # Initialize EasyOCR reader (replace 'en' with your required language(s))
# reader = Reader(['en'])

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using EasyOCR
#     text_detected = reader.readtext(img)

#     # Check if text was detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Process detected text (e.g., join lines, remove non-ASCII characters if needed)
#     text_detected_cleaned = ' '.join([item[1] for item in text_detected])

#     # Join words with spaces
#     if text_detected:
#         try:
#             text_detected_cleaned = ' '.join([line[1] for line in text_detected])
#         except (IndexError, TypeError):
#             text_detected_cleaned = ''
#     else:
#         text_detected_cleaned = ''

#     print(text_detected, "The PRINT OF TEXT DETECTED")

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_pdf_file:
#         # Convert detected text to PDF
#         pdf = FPDF()
#         pdf.add_page()
#         pdf.set_font("Arial", size=12)
#         pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#         pdf.cell(200, 10, txt="", ln=True)
#         # Convert the text to bytes using UTF-8 encoding
#         text_bytes = text_detected_cleaned.encode('utf-8')
#         # Write the text to the PDF using UTF-8 encoding
#         pdf.multi_cell(0, 10, txt=text_bytes.decode('utf-8'))
#         # Output the PDF with UTF-8 encoding
#         pdf.output(tmp_pdf_file.name, "F")
#         tmp_pdf_file.seek(0)

#         # Send the temporary PDF file as an attachment
#         pdf_response = make_response(send_file(tmp_pdf_file.name, as_attachment=True, download_name='recognized_text.pdf'))

#     # Create a Word document with the detected text
#     doc = docx.Document()
#     doc.add_heading('Handwriting Recognition Result', 0)
#     doc.add_paragraph(text_detected_cleaned)

#     # Save the Word document to a temporary file
#     temp_word_file = io.BytesIO()
#     doc.save(temp_word_file)
#     temp_word_file.seek(0)

#     # Send the temporary Word file as an attachment
#     word_response = make_response(send_file(temp_word_file, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', download_name='recognized_text.docx'))

#     # Combine the PDF and Word responses
#     combined_response = make_response(pdf_response.data + b'\n\n' + word_response.data)
#     combined_response.headers.set('Content-Type', 'application/octet-stream')
#     combined_response.headers.set('Content-Disposition', 'attachment', filename='recognized_text.zip')

#     return combined_response

# if __name__ == '__main__':
#     app.run(host='0.0.0.0', port=5000, debug=True)

                            #WOrking




# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# from easyocr import Reader
# import os
# import docx  # Import the python-docx library

# os.environ['CUDA_VISIBLE_DEVICES'] = '-1'
# app = Flask(__name__)

# # Configure paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# # Initialize EasyOCR reader (replace 'en' with your required language(s))
# reader = Reader(['en'])

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using EasyOCR
#     text_detected = reader.readtext(img)

#     # Check if text was detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Process detected text (e.g., join lines, remove non-ASCII characters if needed)
#     text_detected_cleaned = ' '.join([item[1] for item in text_detected])

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_pdf_file:
#         # Convert detected text to PDF
#         pdf = FPDF()
#         pdf.add_page()
#         pdf.set_font("Arial", size=12)
#         pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#         pdf.cell(200, 10, txt="", ln=True)
#         # Convert the text to bytes using UTF-8 encoding
#         text_bytes = text_detected_cleaned.encode('utf-8')
#         # Write the text to the PDF using UTF-8 encoding
#         pdf.multi_cell(0, 10, txt=text_bytes.decode('utf-8'))
#         # Output the PDF with UTF-8 encoding
#         pdf.output(tmp_pdf_file.name, "F")
#         tmp_pdf_file.seek(0)

#         # Create a Word document with the detected text
#         doc = docx.Document()
#         doc.add_heading('Handwriting Recognition Result', 0)
#         doc.add_paragraph(text_detected_cleaned)

#         # Save the Word document to a temporary file
#         temp_word_file = io.BytesIO()
#         doc.save(temp_word_file)
#         temp_word_file.seek(0)

#     # Create a ZIP file containing both PDF and Word files
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_zip_file:
#         # Add PDF and Word files to the ZIP
#         with zipfile.ZipFile(tmp_zip_file, 'w') as zipf:
#             zipf.write(tmp_pdf_file.name, 'recognized_text.pdf')
#             zipf.write(temp_word_file, 'recognized_text.docx')

#     # Send the ZIP file as an attachment
#     zip_response = make_response(send_file(tmp_zip_file.name, as_attachment=True, mimetype='application/zip', download_name='recognized_text.zip'))

#     return zip_response

# if __name__ == '__main__':
#     app.run(host='0.0.0.0', port=5000, debug=True)




# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# import pytesseract

# app = Flask(__name__)

# # Configure pytesseract path to Tesseract executable
# # Update the path below to where Tesseract is installed on your system
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using pytesseract
#     text_detected = pytesseract.image_to_string(img)

#     # Check if text was detected
#     if not text_detected.strip():
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Remove non-ASCII characters from the text
#     text_detected_cleaned = ''.join(char for char in text_detected if ord(char) < 128)

#     # Convert detected text to PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#     pdf.cell(200, 10, txt="", ln=True)
#     pdf.multi_cell(0, 10, txt=text_detected_cleaned)

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         pdf.output(tmp_file.name)
#         tmp_file.seek(0)

#     # Send the temporary file as an attachment
#     response = make_response(send_file(tmp_file.name, as_attachment=True))
#     response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#     return response

# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)



# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# from easyocr import Reader
# import os
# os.environ['CUDA_VISIBLE_DEVICES'] = '-1'


# app = Flask(__name__)

# # Configure paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# # Initialize EasyOCR reader (replace 'en' with your required language(s))
# reader = Reader(['en'])

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using EasyOCR
#     text_detected = reader.readtext(img)

#     # Check if text was detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Process detected text (e.g., join lines, remove non-ASCII characters if needed)
#     text_detected_cleaned = ' '.join([line[1] for line in text_detected])  # Join words with spaces

#     # Convert detected text to PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#     pdf.cell(200, 10, txt="", ln=True)
#     pdf.multi_cell(0, 10, txt=text_detected_cleaned)

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         pdf.output(tmp_file.name)
#         tmp_file.seek(0)

#     # Send the temporary file as an attachment
#     response = make_response(send_file(tmp_file.name, as_attachment=True))
#     response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#     return response


# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)






#  Working Code without Preprocessing:
# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# from easyocr import Reader
# import os

# os.environ['CUDA_VISIBLE_DEVICES'] = '-1'
# app = Flask(__name__)

# # Configure paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# # Initialize EasyOCR reader (replace 'en' with your required language(s))
# reader = Reader(['en'])

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using EasyOCR
#     text_detected = reader.readtext(img)

#     # Check if text was detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Process detected text (e.g., join lines, remove non-ASCII characters if needed)
#     text_detected_cleaned = ' '.join([item[1] for item in text_detected])

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         # Convert detected text to PDF
#         pdf = FPDF()
#         pdf.add_page()
#         pdf.set_font("Arial", size=12)
#         pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#         pdf.cell(200, 10, txt="", ln=True)
#         # Convert the text to bytes using UTF-8 encoding
#         text_bytes = text_detected_cleaned.encode('utf-8')
#         # Write the text to the PDF using UTF-8 encoding
#         pdf.multi_cell(0, 10, txt=text_bytes.decode('utf-8'))
#         # Output the PDF with UTF-8 encoding
#         pdf.output(tmp_file.name)

#         # Send the temporary file as an attachment
#         response = make_response(send_file(tmp_file.name, as_attachment=True, mimetype='application/pdf'))

#         return response


# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)






















                #Using Grascale in easyocr
# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# from easyocr import Reader
# import os

# os.environ['CUDA_VISIBLE_DEVICES'] = '-1'
# app = Flask(__name__)

# # Configure paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# # Initialize EasyOCR reader (replace 'en' with your required language(s))
# reader = Reader(['en'])

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Read the image file as numpy array
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)

#     # Convert numpy array to OpenCV format
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Preprocessing: Convert image to grayscale
#     gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
#     # Perform OCR using EasyOCR on preprocessed image
#     result = reader.readtext(gray)

#     # Check if text was detected
#     if not result:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Process detected text
#     text_detected_cleaned = ''

#     for detection in result:
#         text = detection[1]
#         # If there's already text appended, add a space before appending new text
#         if text_detected_cleaned:
#             text_detected_cleaned += ' '
#         text_detected_cleaned += text

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         # Convert detected text to PDF
#         pdf = FPDF()
#         pdf.add_page()
#         pdf.set_font("Arial", size=12)
#         pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#         pdf.cell(200, 10, txt="", ln=True)
#         # Convert the text to bytes using UTF-8 encoding
#         text_bytes = text_detected_cleaned.encode('utf-8')
#         # Write the text to the PDF using UTF-8 encoding
#         pdf.multi_cell(0, 10, txt=text_bytes.decode('utf-8'))
#         # Output the PDF with UTF-8 encoding
#         pdf.output(tmp_file.name)

#         # Send the temporary file as an attachment
#         response = make_response(send_file(tmp_file.name, as_attachment=True, mimetype='application/pdf'))

#         return response



# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)





# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# import pytesseract

# app = Flask(__name__)

# # Configure pytesseract path to Tesseract executable
# # Update the path below to where Tesseract is installed on your system
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# # Define paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using pytesseract
#     text_detected = pytesseract.image_to_string(img)

#     # Check if text was detected
#     if not text_detected.strip():
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Remove non-ASCII characters from the text
#     text_detected_cleaned = ''.join(char for char in text_detected if ord(char) < 128)

#     # Convert detected text to PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#     pdf.cell(200, 10, txt="", ln=True)
#     pdf.multi_cell(0, 10, txt=text_detected_cleaned)

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         pdf.output(tmp_file.name)
#         tmp_file.seek(0)

#     # Send the temporary file as an attachment
#     response = make_response(send_file(tmp_file.name, as_attachment=True))
#     response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#     return response

# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)









        #  Download As Word with Unreadable Error



# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# import pytesseract
# import docx
# from docx import Document
# import locale
# from docx.shared import Inches

# app = Flask(__name__)

# # Configure pytesseract path to Tesseract executable
# # Update the path below to where Tesseract is installed on your system
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# # Define paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using pytesseract
#     text_detected = pytesseract.image_to_string(img)

#     # Check if text was detected
#     if not text_detected.strip():
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Remove non-ASCII characters from the text
#     text_detected_cleaned = ''.join(char for char in text_detected if ord(char) < 128)

#     # Convert detected text to PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#     pdf.cell(200, 10, txt="", ln=True)
#     pdf.multi_cell(0, 10, txt=text_detected_cleaned)

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         pdf.output(tmp_file.name)
#         tmp_file.seek(0)

#     # Send the temporary file as an attachment
#     response = make_response(send_file(tmp_file.name, as_attachment=True))
#     response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#     return response

# @app.route('/download_text', methods=['POST'])
# def download_text():
#     # Get the text content from the request
#     text_content = request.form.get('text_content')

#     # Create a Word document with the text content
#     doc = Document()
#     doc.add_heading('Handwriting Recognition Result', 0)
#     doc.add_paragraph(text_content)

#     # Create a temporary file to save the Word document
#     with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
#         doc.save(tmp_file.name)
#         tmp_file.seek(0)
#     # Create a temporary file to save the Word document
#     with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp_file:
#         doc.save(tmp_file.name)
#         tmp_file.seek(0)

#         # Send the temporary file as an attachment
#         response = make_response(send_file(tmp_file.name, as_attachment=True))
#         response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.doc'
#         return response


# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)



# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# import pytesseract
# from docx import Document
# import locale
# from docx.shared import Inches

# app = Flask(__name__)

# # Configure pytesseract path to Tesseract executable
# # Update the path below to where Tesseract is installed on your system
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# # Define paths to CSS and JavaScript files
# CSS_PATH = "../static/style.css"
# SCRIPT_PATH = "../static/script.js"

# @app.route('/')
# def index():
#     return render_template('index.html', css_path=CSS_PATH, script_path=SCRIPT_PATH)

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using pytesseract
#     text_detected = pytesseract.image_to_string(img)

#     # Check if text was detected
#     if not text_detected.strip():
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Remove non-ASCII characters from the text
#     text_detected_cleaned = ''.join(char for char in text_detected if ord(char) < 128)

#     # Convert detected text to PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#     pdf.cell(200, 10, txt="", ln=True)
#     pdf.multi_cell(0, 10, txt=text_detected_cleaned)

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         pdf.output(tmp_file.name)
#         tmp_file.seek(0)

#     # Send the temporary file as an attachment
#     response = make_response(send_file(tmp_file.name, as_attachment=True))
#     response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#     return response

# @app.route('/download_text', methods=['POST'])
# def download_text():
#     # Get the text content from the request
#     text_content = request.form.get('text_content')

#     # Create a temporary text file
#     with tempfile.NamedTemporaryFile(delete=False, mode='w', encoding='utf-8') as tmp_file:
#         tmp_file.write(text_content)
#         tmp_file.seek(0)

#         # Send the temporary file as an attachment
#         response = make_response(send_file(tmp_file.name, as_attachment=True))
#         response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.txt'
#         return response
    

# if __name__ == '__main__':
#     # Run Flask app on all network interfaces
#     app.run(host='0.0.0.0', port=5000, debug=True)









# easyocr word file download

# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# import easyocr
# from docx import Document

# app = Flask(__name__)

# # Define paths to CSS and JavaScript files
# CSS_PATH = "../style.css"
# SCRIPT_PATH = "../script.js"

# # Create an EasyOCR reader
# reader = easyocr.Reader(['en'])

# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR
#     text_detected = reader.readtext(img)

#     # Check if text was detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Join the detected text
#     text_detected_cleaned = ' '.join([text[1] for text in text_detected])

#     # Convert detected text to PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#     pdf.cell(200, 10, txt="", ln=True)
#     pdf.multi_cell(0, 10, txt=text_detected_cleaned)

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         pdf.output(tmp_file.name)
#         tmp_file.seek(0)

#         # Send the temporary file as an attachment
#         response = make_response(send_file(tmp_file.name, as_attachment=True))
#         response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#         return response

# @app.route('/download_text', methods=['POST'])
# def download_text():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR
#     text_detected = reader.readtext(img)

#     # Check if text was detected
#     if not text_detected:
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Join the detected text
#     text_detected_cleaned = ' '.join([text[1] for text in text_detected])

#     # Create a Word document with the text content
#     doc = Document()
#     doc.add_heading('Handwriting Recognition Result', 0)
#     doc.add_paragraph(text_detected_cleaned)

#     # Create a temporary file to save the Word document
#     with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
#         doc.save(tmp_file.name)
#         tmp_file.seek(0)

#         # Send the temporary file as an attachment
#         response = make_response(send_file(tmp_file.name, as_attachment=True))
#         response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.docx'
#         return response

# if __name__ == '__main__':
#     import os
#     port = int(os.environ.get('PORT', 5000))
#     app.run(host='0.0.0.0', port=port)



# direct pytesseract library use instead of EXE


# from flask import Flask, render_template, request, jsonify, send_file, make_response
# import cv2
# import numpy as np
# import io
# from fpdf import FPDF
# import tempfile
# import pytesseract
# from docx import Document

# app = Flask(__name__)

# # Define paths to CSS and JavaScript files
# CSS_PATH = "../style.css"
# SCRIPT_PATH = "../script.js"

# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route('/recognize_handwriting', methods=['POST'])
# def recognize_handwriting():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using pytesseract
#     text_detected = pytesseract.image_to_string(img)

#     # Check if text was detected
#     if not text_detected.strip():
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Remove non-ASCII characters from the text
#     text_detected_cleaned = ''.join(char for char in text_detected if ord(char) < 128)

#     # Convert detected text to PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
#     pdf.cell(200, 10, txt="", ln=True)
#     pdf.multi_cell(0, 10, txt=text_detected_cleaned)

#     # Create a temporary file to save the PDF
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         pdf.output(tmp_file.name)
#         tmp_file.seek(0)

#         # Send the temporary file as an attachment
#         response = make_response(send_file(tmp_file.name, as_attachment=True))
#         response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
#         return response

# @app.route('/download_text', methods=['POST'])
# def download_text():
#     # Get the image file from the request
#     image_file = request.files['image']

#     # Convert the image data to OpenCV format
#     img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
#     img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

#     # Check if the image was loaded successfully
#     if img is None:
#         return jsonify({'error': 'Unable to read the image'}), 500

#     # Perform OCR using pytesseract
#     text_detected = pytesseract.image_to_string(img)

#     # Check if text was detected
#     if not text_detected.strip():
#         return jsonify({'error': 'No text detected in the image'}), 400

#     # Remove non-ASCII characters from the text
#     text_detected_cleaned = ''.join(char for char in text_detected if ord(char) < 128)

#     # Create a Word document with the text content
#     doc = Document()
#     doc.add_heading('Handwriting Recognition Result', 0)
#     doc.add_paragraph(text_detected_cleaned)

#     # Create a temporary file to save the Word document
#     with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
#         doc.save(tmp_file.name)
#         tmp_file.seek(0)

#         # Send the temporary file as an attachment
#         response = make_response(send_file(tmp_file.name, as_attachment=True))
#         response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.docx'
#         return response

# if __name__ == '__main__':
#      app.run(host='0.0.0.0', port=5000, debug=True)





























# MOST ACCURATE WORKING CODE

from flask import Flask, render_template, request, jsonify, send_file, make_response
import cv2
import numpy as np
import io
from fpdf import FPDF
import tempfile
import pytesseract
import docx
from docx import Document
import locale

app = Flask(__name__)

# Configure pytesseract path to Tesseract executable
# Update the path below to where Tesseract is installed on your system
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Define paths to CSS and JavaScript files
CSS_PATH = "../style.css"
SCRIPT_PATH = "../script.js"

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/recognize_handwriting', methods=['POST'])
def recognize_handwriting():
    # Get the image file from the request
    image_file = request.files['image']

    # Convert the image data to OpenCV format
    img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
    img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

    # Check if the image was loaded successfully
    if img is None:
        return jsonify({'error': 'Unable to read the image'}), 500

    # Perform OCR using pytesseract
    text_detected = pytesseract.image_to_string(img)

    # Check if text was detected
    if not text_detected.strip():
        return jsonify({'error': 'No text detected in the image'}), 400

    # Remove non-ASCII characters from the text
    text_detected_cleaned = ''.join(char for char in text_detected if ord(char) < 128)

    # Convert detected text to PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Handwriting Recognition Result", ln=True, align="C")
    pdf.cell(200, 10, txt="", ln=True)
    pdf.multi_cell(0, 10, txt=text_detected_cleaned)

    # Create a temporary file to save the PDF
    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        pdf.output(tmp_file.name)
        tmp_file.seek(0)

        # Send the temporary file as an attachment
        response = make_response(send_file(tmp_file.name, as_attachment=True))
        response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.pdf'
        return response

@app.route('/download_text', methods=['POST'])
def download_text():
    # Get the image file from the request
    image_file = request.files['image']

    # Convert the image data to OpenCV format
    img_array = np.frombuffer(image_file.read(), dtype=np.uint8)
    img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

    # Check if the image was loaded successfully
    if img is None:
        return jsonify({'error': 'Unable to read the image'}), 500

    # Perform OCR using pytesseract
    text_detected = pytesseract.image_to_string(img)

    # Check if text was detected
    if not text_detected.strip():
        return jsonify({'error': 'No text detected in the image'}), 400

    # Remove non-ASCII characters from the text
    text_detected_cleaned = ''.join(char for char in text_detected if ord(char) < 128)

    # Create a Word document with the text content
    doc = Document()
    doc.add_heading('Handwriting Recognition Result', 0)
    doc.add_paragraph(text_detected_cleaned)

    # Create a temporary file to save the Word document
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        doc.save(tmp_file.name)
        tmp_file.seek(0)

        # Send the temporary file as an attachment
        response = make_response(send_file(tmp_file.name, as_attachment=True))
        response.headers['Content-Disposition'] = 'attachment; filename=recognized_text.docx'
        return response

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)